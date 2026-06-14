#!/usr/bin/env python3
"""
build_rationale_docx.py — format a Chinese 《修改理由》 Markdown into Word, in the
house format the user requires for these explainer docs.

Usage:
    python build_rationale_docx.py 修改理由.md manuscript_修改理由.docx

Input Markdown conventions:
    # ...      -> title          (13 pt, bold)
    ## ...     -> subheading     (11 pt, bold)
    ### ...    -> subheading     (11 pt, bold)
    anything else (incl. "1. **改动**：…" / "**理由**：…") -> body (11 pt)
    inline **bold** is honoured everywhere.

House format enforced (all via DIRECT formatting — no heading/semantic styles, "flat"):
    - hierarchy distinguished only by font size + bold
    - line spacing 1.15; space-after 0.5 line
    - body justified (两端对齐)
    - Chinese in 宋体 (rFonts w:eastAsia); English/digits in Times New Roman (ascii/hAnsi/cs)
    - A4 page, 1-inch (25.4 mm) margins

Depends on python-docx.
"""
import re
import sys
import argparse
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TITLE_PT = 13
SUB_PT = 11
BODY_PT = 11
CJK_FONT = "宋体"
LATIN_FONT = "Times New Roman"


def _set_fonts(run):
    rpr = run._element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    for k in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(k), LATIN_FONT)
    rf.set(qn("w:eastAsia"), CJK_FONT)


def _fmt_para(p):
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pPr = p._p.get_or_add_pPr()
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        pPr.append(sp)
    sp.set(qn("w:line"), "276")        # 1.15 * 240
    sp.set(qn("w:lineRule"), "auto")
    sp.set(qn("w:afterLines"), "50")   # 0.5 line
    sp.set(qn("w:after"), "0")
    sp.set(qn("w:beforeLines"), "0")
    sp.set(qn("w:before"), "0")


def _add_runs(p, text, size, bold0=False):
    bold = bold0
    for part in re.split(r"(\*\*)", text):
        if part == "**":
            bold = not bold
            continue
        if not part:
            continue
        r = p.add_run(part)
        r.bold = bold
        r.font.size = Pt(size)
        _set_fonts(r)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("src", help="input rationale Markdown")
    ap.add_argument("out", help="output .docx")
    args = ap.parse_args()

    doc = Document()
    s = doc.sections[0]
    s.page_width, s.page_height = Mm(210), Mm(297)
    s.top_margin = s.bottom_margin = Mm(25.4)
    s.left_margin = s.right_margin = Mm(25.4)

    # Normal-style defaults so anything not explicitly set is still consistent
    st = doc.styles["Normal"]
    st.font.name = LATIN_FONT
    st.font.size = Pt(BODY_PT)
    rf = st.element.get_or_add_rPr().get_or_add_rFonts()
    for k in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(k), LATIN_FONT)
    rf.set(qn("w:eastAsia"), CJK_FONT)

    with open(args.src, encoding="utf-8") as fh:
        for raw in fh:
            ln = raw.rstrip("\n")
            if not ln.strip():
                continue
            if ln.startswith("# "):
                p = doc.add_paragraph(); _fmt_para(p); _add_runs(p, ln[2:].strip(), TITLE_PT, True)
            elif ln.startswith("## "):
                p = doc.add_paragraph(); _fmt_para(p); _add_runs(p, ln[3:].strip(), SUB_PT, True)
            elif ln.startswith("### "):
                p = doc.add_paragraph(); _fmt_para(p); _add_runs(p, ln[4:].strip(), SUB_PT, True)
            else:
                p = doc.add_paragraph(); _fmt_para(p); _add_runs(p, ln.strip(), BODY_PT, False)

    doc.save(args.out)
    print("wrote", args.out)


if __name__ == "__main__":
    main()
